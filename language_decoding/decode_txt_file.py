from googletrans import Translator
import pandas as pd
import unicodedata
from docx import Document
import argparse
### Story data:
#http://www.mikelockett.com/storytelling?lang=spanish    


def data_preprocessing(data):
    data_string = ""
    for i in data:
        #print(i)
        data_string = data_string+i
    cleaned_data = data_string.replace("\n"," ")
    word_list = cleaned_data.split(" ")
    return word_list, cleaned_data

def translate(args, word_list):
    translator = Translator()
    src = args.Source_language
    dest = args.Destination_language
    destination_word_list = []
    print("Source Language: {} -- Destination Language: {}".format(src,dest))
    #print(word_list)
    for i in word_list:
        #try:
            #translator = Translator()
            #############################################################################
        #print(i)
        if i ==" ":
            pass

        #print(i)
        translator = Translator()
        word_translation = translator.translate(i, src= src, dest= dest)

        #except:
        #    translator = Translator(i)
        #    word_translation = translator.translate(i)
        destination_word_list.append(word_translation.text)
            ###############################################################################

    return destination_word_list


def main(args):
    try:
        data_name = args.Title
        if data_name.endswith(".txt"):
            pass
        else:
            data_name = data_name+".txt"
        data = open(data_name,"r")

        word_list, cleaned_data = data_preprocessing(data)
        destination_word_list = translate(args, word_list)

        if args.Output == "Table":
            print("Writing data to table...")
            data_dic = {"Original_data": word_list, "Translated_data": destination_word_list}
            data_frame = pd.DataFrame(data_dic)

            print("Please insert a name for your data-table: \n")
            name = input()
            data_frame.to_csv(name)
            print("Table is saved as {}.csv".format(name))

        elif args.Output == "Doc":
            final_string = ""
            for i in range(len(destination_word_list)):
                if i == 0:
                    final_string = final_string + destination_word_list[i]
                else:
                    final_string = final_string+" "+ destination_word_list[i]
            print(final_string)
            document = Document()
            data_name = data_name.split(".txt")
            document.add_heading(data_name[0], 0)

            document.add_paragraph(cleaned_data)

            document.add_paragraph(final_string)
            document.add_page_break()

            print("Data got written to the document! \nPlease insert a name to save the document:")
            name = input()
            if name.endswith(".docx"):
                document.save(name)
            else:
                name = name+".docx"
                document.save(name)
            print("Document is saved as {}".format(name))

        else: 
            print("Error while saving translation!")

    finally:
        print("----")

if __name__ == "__main__":
    parser = argparse.ArgumentParser()
    parser.add_argument("-t", "--Title", type = str, help = "The name of the txt file that you want to decode word by word!")
    parser.add_argument("-src", "--Source_language",type = str, help = "Source Language")
    parser.add_argument("-dest", "--Destination_language",type = str, help = "Destination Language")
    parser.add_argument("-out", "--Output",type = str,choices = ["Doc", "Table"], help = "What should the output format be")

    args = parser.parse_args()
    main(args)
